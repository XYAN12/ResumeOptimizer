from __future__ import annotations

import base64
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader, PdfWriter
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.pdfmetrics import registerFont
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.core.config import get_settings
from app.core.exceptions import ExportError
from app.models.domain import RewriteResult


class ExportService:
    def export(
        self,
        session_id: str,
        rewrite: RewriteResult,
        export_format: str,
        original_file_content: bytes | None = None,
        original_filename: str | None = None,
    ) -> tuple[str, str]:
        filename = f"optimized_resume_{session_id}.{export_format}"
        if export_format == "md":
            payload = rewrite.markdown.encode("utf-8")
        elif export_format == "docx":
            payload = self._to_docx(rewrite, original_file_content, original_filename)
        elif export_format == "pdf":
            payload = self._to_pdf(rewrite, original_file_content, original_filename)
        else:
            raise ExportError(f"Unsupported export format: {export_format}")

        settings = get_settings()
        export_dir = Path(settings.export_dir)
        export_dir.mkdir(parents=True, exist_ok=True)
        (export_dir / filename).write_bytes(payload)
        return filename, base64.b64encode(payload).decode("utf-8")

    def _to_docx(
        self,
        rewrite: RewriteResult,
        original_file_content: bytes | None,
        original_filename: str | None,
    ) -> bytes:
        # Export from rewritten preview structure directly for consistent output.
        document = Document()
        for section in rewrite.sections:
            if section.title.lower() == "header":
                for item in section.items:
                    document.add_paragraph(item)
                continue

            document.add_heading(section.title, level=2)
            for item in section.items:
                document.add_paragraph(item, style="List Bullet")
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _to_template_docx(self, rewrite: RewriteResult, original_docx_bytes: bytes) -> bytes:
        document = Document(BytesIO(original_docx_bytes))
        ranges = self._locate_docx_section_ranges(document, [section.title for section in rewrite.sections])

        for section in rewrite.sections:
            paragraph_range = ranges.get(section.title)
            if not paragraph_range:
                continue
            start_idx, end_idx = paragraph_range
            section_paragraphs = document.paragraphs[start_idx:end_idx]
            if not section_paragraphs:
                continue
            self._rewrite_docx_section_paragraphs(section.title, section.items, section_paragraphs)

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    def _locate_docx_section_ranges(
        self,
        document: Document,
        section_titles: list[str],
    ) -> dict[str, tuple[int, int]]:
        paragraphs = document.paragraphs
        heading_candidates = [title for title in section_titles if title.lower() != "header"]
        heading_positions: dict[str, int] = {}

        for idx, paragraph in enumerate(paragraphs):
            normalized = self._normalize_title(paragraph.text)
            for title in heading_candidates:
                if title in heading_positions:
                    continue
                if normalized == self._normalize_title(title):
                    heading_positions[title] = idx

        ordered_headings = sorted(heading_positions.items(), key=lambda item: item[1])
        ranges: dict[str, tuple[int, int]] = {}

        if "Header" in section_titles:
            first_heading_index = ordered_headings[0][1] if ordered_headings else len(paragraphs)
            ranges["Header"] = (0, first_heading_index)

        for index, (title, start_idx) in enumerate(ordered_headings):
            next_start = ordered_headings[index + 1][1] if index + 1 < len(ordered_headings) else len(paragraphs)
            ranges[title] = (start_idx + 1, next_start)
        return ranges

    def _rewrite_docx_section_paragraphs(self, title: str, optimized_items: list[str], paragraphs: list) -> None:
        target_paragraphs = [paragraph for paragraph in paragraphs if paragraph.text.strip()]
        if not target_paragraphs:
            target_paragraphs = paragraphs
        if not target_paragraphs:
            return

        fitted_items = self._fit_items_to_slots(optimized_items, len(target_paragraphs))
        for idx, paragraph in enumerate(target_paragraphs):
            prefix = "" if title.lower() == "header" else "• "
            paragraph.text = f"{prefix}{fitted_items[idx]}"

    def _to_pdf(
        self,
        rewrite: RewriteResult,
        original_file_content: bytes | None,
        original_filename: str | None,
    ) -> bytes:
        self._register_pdf_fonts()
        # Export from rewritten preview structure directly for consistent output.
        return self._to_plain_pdf(rewrite)

    def _to_plain_pdf(self, rewrite: RewriteResult) -> bytes:
        buffer = BytesIO()
        document = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=40,
            rightMargin=40,
            topMargin=40,
            bottomMargin=40,
        )
        stylesheet = getSampleStyleSheet()
        base_style = ParagraphStyle(
            "ResumeBase",
            parent=stylesheet["BodyText"],
            fontName="STSong-Light",
            fontSize=10.5,
            leading=15,
            alignment=TA_LEFT,
            wordWrap="CJK",
        )
        heading_style = ParagraphStyle(
            "ResumeHeading",
            parent=base_style,
            fontSize=14,
            leading=18,
            spaceBefore=8,
            spaceAfter=8,
        )
        story = []
        for section in rewrite.sections:
            if section.title.lower() == "header":
                for item in section.items:
                    story.append(Paragraph(self._escape_pdf_text(item), base_style))
                story.append(Spacer(1, 10))
                continue

            story.append(Paragraph(self._escape_pdf_text(section.title), heading_style))
            for item in section.items:
                story.append(Paragraph(self._escape_pdf_text(f"• {item}"), base_style))
            story.append(Spacer(1, 8))

        document.build(story)
        return buffer.getvalue()

    def _to_template_pdf(self, rewrite: RewriteResult, original_pdf_bytes: bytes) -> bytes:
        base_reader = PdfReader(BytesIO(original_pdf_bytes))
        overlay_buffer = BytesIO()
        overlay_canvas = canvas.Canvas(overlay_buffer)

        for page_index, page in enumerate(base_reader.pages, start=1):
            width = float(page.mediabox.width)
            height = float(page.mediabox.height)
            overlay_canvas.setPageSize((width, height))
            overlay_canvas.setFont("STSong-Light", 10.5)

            page_sections = [
                section
                for section in rewrite.sections
                if section.layout and section.layout.page_number == page_index
            ]
            for section in page_sections:
                self._draw_section_on_template(overlay_canvas, section, height)

            overlay_canvas.showPage()

        overlay_canvas.save()
        overlay_reader = PdfReader(BytesIO(overlay_buffer.getvalue()))
        writer = PdfWriter()

        for page_index, page in enumerate(base_reader.pages):
            merged_page = page
            if page_index < len(overlay_reader.pages):
                merged_page.merge_page(overlay_reader.pages[page_index])
            writer.add_page(merged_page)

        output = BytesIO()
        writer.write(output)
        return output.getvalue()

    def _draw_section_on_template(self, pdf: canvas.Canvas, section, page_height: float) -> None:
        layout = section.layout
        if not layout:
            return

        x = layout.x or 40.0
        width = layout.width or 500.0
        body_font_size = layout.body_font_size or section.layout.title_font_size or 10.5
        start_y = (layout.y_top or page_height - 60) - ((layout.title_font_size or 12.0) + 8)
        bottom_limit = layout.y_bottom or 48.0
        leading = max(body_font_size + 4, 13.5)

        pdf.setFillColorRGB(0, 0, 0)
        pdf.setFont("STSong-Light", body_font_size)

        y = start_y
        bullet_prefix = "" if section.title.lower() == "header" else "• "
        fitted_items = self._fit_items_to_slots(section.items, max(len(section.supporting_facts), 1))
        for item in fitted_items:
            wrapped_lines = self._wrap_for_template(
                f"{bullet_prefix}{item}",
                width=width,
                font_size=body_font_size,
            )
            for line in wrapped_lines:
                if y <= bottom_limit:
                    return
                pdf.drawString(x, y, line)
                y -= leading

    def _wrap_for_template(self, text: str, width: float, font_size: float) -> list[str]:
        approximate_char_capacity = max(int(width / max(font_size, 1.0)), 8)
        lines: list[str] = []
        current = ""

        for char in text:
            candidate = f"{current}{char}"
            if len(candidate) <= approximate_char_capacity:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = char
        if current:
            lines.append(current)
        return lines or [text]

    def _register_pdf_fonts(self) -> None:
        try:
            registerFont(UnicodeCIDFont("STSong-Light"))
        except Exception:
            pass

    def _escape_pdf_text(self, text: str) -> str:
        return (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br/>")
        )

    def _normalize_title(self, text: str) -> str:
        return re.sub(r"[\W_]+", "", text.lower(), flags=re.UNICODE).strip()

    def _fit_items_to_slots(self, items: list[str], slot_count: int) -> list[str]:
        if slot_count <= 0:
            return []
        normalized = [item.strip() for item in items if item and item.strip()]
        if not normalized:
            return [""] * slot_count
        if len(normalized) == slot_count:
            return normalized
        if len(normalized) > slot_count:
            merged: list[str] = []
            avg = len(normalized) / slot_count
            start = 0
            for index in range(slot_count):
                end = round((index + 1) * avg)
                chunk = normalized[start:end] if end > start else [normalized[start]]
                merged.append("；".join(part for part in chunk if part))
                start = end
            return merged

        # len(normalized) < slot_count: duplicate semantic flow by splitting long lines first.
        expanded = list(normalized)
        while len(expanded) < slot_count:
            longest_index = max(range(len(expanded)), key=lambda idx: len(expanded[idx]))
            longest = expanded[longest_index]
            split_index = self._smart_split_index(longest)
            if split_index <= 0 or split_index >= len(longest):
                expanded.append("")
            else:
                first = longest[:split_index].strip(" ，,;；")
                second = longest[split_index:].strip(" ，,;；")
                expanded[longest_index] = first or longest
                expanded.insert(longest_index + 1, second or "")
        return expanded[:slot_count]

    def _smart_split_index(self, text: str) -> int:
        if len(text) < 20:
            return -1
        middle = len(text) // 2
        delimiters = "，,；;。 "
        for offset in range(0, len(text) // 2):
            right = middle + offset
            left = middle - offset
            if right < len(text) and text[right] in delimiters:
                return right + 1
            if left >= 0 and text[left] in delimiters:
                return left + 1
        return middle
