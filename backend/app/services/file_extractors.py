from __future__ import annotations

from collections import Counter
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Length
from pypdf import PdfReader

from app.core.exceptions import FileExtractionError
from app.models.domain import DocumentTheme, ExtractedResumeDocument, LayoutLine


def extract_resume_document(filename: str, content: bytes) -> ExtractedResumeDocument:
    suffix = Path(filename).suffix.lower()

    try:
        if suffix in {".txt", ".md"}:
            text = content.decode("utf-8")
            return ExtractedResumeDocument(
                text=text,
                source_format=suffix.lstrip("."),
                theme=DocumentTheme(source_format=suffix.lstrip(".")),
                filename=filename,
                content=content,
            )
        if suffix == ".pdf":
            return _extract_pdf_document(filename, content)
        if suffix == ".docx":
            return _extract_docx_document(filename, content)
    except Exception as exc:  # pragma: no cover - defensive wrapper
        raise FileExtractionError(f"Failed to extract text from {filename}") from exc

    raise FileExtractionError(f"Unsupported file type: {suffix}")


def _extract_pdf_document(filename: str, content: bytes) -> ExtractedResumeDocument:
    reader = PdfReader(BytesIO(content))
    layout_lines: list[LayoutLine] = []
    page_width = None
    page_height = None

    for page_number, page in enumerate(reader.pages, start=1):
        page_width = float(page.mediabox.width)
        page_height = float(page.mediabox.height)
        spans: list[LayoutLine] = []

        def visitor_text(text, _cm, tm, font_dict, font_size) -> None:
            cleaned = " ".join(text.replace("\x00", "").split())
            if not cleaned:
                return

            font_name = None
            if font_dict:
                font_name = str(font_dict.get("/BaseFont", "")).lstrip("/")

            spans.append(
                LayoutLine(
                    text=cleaned,
                    page_number=page_number,
                    x=float(tm[4]) if len(tm) > 4 else None,
                    y=float(tm[5]) if len(tm) > 5 else None,
                    font_size=float(font_size) if font_size else None,
                    font_name=font_name,
                )
            )

        page.extract_text(visitor_text=visitor_text)
        layout_lines.extend(_group_pdf_spans_into_lines(spans))

    ordered_lines = sorted(
        layout_lines,
        key=lambda line: (line.page_number, -(line.y or 0.0), line.x or 0.0),
    )
    text = "\n".join(line.text for line in ordered_lines).strip()
    theme = _infer_pdf_theme(ordered_lines, page_width, page_height)

    return ExtractedResumeDocument(
        text=text,
        source_format="pdf",
        lines=ordered_lines,
        theme=theme,
        filename=filename,
        content=content,
    )


def _extract_docx_document(filename: str, content: bytes) -> ExtractedResumeDocument:
    document = Document(BytesIO(content))
    lines: list[LayoutLine] = []
    font_sizes: list[float] = []
    font_names: list[str] = []

    for index, paragraph in enumerate(document.paragraphs):
        text = " ".join(paragraph.text.split()).strip()
        if not text:
            continue

        line_font_size = _docx_paragraph_font_size(paragraph)
        line_font_name = _docx_paragraph_font_name(paragraph)
        if line_font_size:
            font_sizes.append(line_font_size)
        if line_font_name:
            font_names.append(line_font_name)

        lines.append(
            LayoutLine(
                text=text,
                page_number=1,
                x=0.0,
                y=float(10_000 - index),
                font_size=line_font_size,
                font_name=line_font_name,
            )
        )

    text = "\n".join(line.text for line in lines).strip()
    body_font_size = Counter(round(size, 1) for size in font_sizes).most_common(1)[0][0] if font_sizes else None
    body_font_name = Counter(font_names).most_common(1)[0][0] if font_names else None

    return ExtractedResumeDocument(
        text=text,
        source_format="docx",
        lines=lines,
        theme=DocumentTheme(
            source_format="docx",
            body_font_name=body_font_name,
            heading_font_name=body_font_name,
            body_font_size=body_font_size,
            heading_font_size=body_font_size,
            preserve_original_docx_look=True,
        ),
        filename=filename,
        content=content,
    )


def _group_pdf_spans_into_lines(spans: list[LayoutLine]) -> list[LayoutLine]:
    if not spans:
        return []

    ordered = sorted(spans, key=lambda span: (-(span.y or 0.0), span.x or 0.0))
    grouped: list[list[LayoutLine]] = []

    for span in ordered:
        if not grouped:
            grouped.append([span])
            continue

        last_group = grouped[-1]
        last_y = last_group[0].y or 0.0
        current_y = span.y or 0.0
        if abs(last_y - current_y) <= 2.5:
            last_group.append(span)
        else:
            grouped.append([span])

    lines: list[LayoutLine] = []
    for group in grouped:
        line_spans = sorted(group, key=lambda span: span.x or 0.0)
        line_text = " ".join(span.text for span in line_spans).strip()
        if not line_text:
            continue

        font_sizes = [span.font_size for span in line_spans if span.font_size]
        font_names = [span.font_name for span in line_spans if span.font_name]
        lines.append(
            LayoutLine(
                text=line_text,
                page_number=line_spans[0].page_number,
                x=min((span.x or 0.0) for span in line_spans),
                y=max((span.y or 0.0) for span in line_spans),
                font_size=max(font_sizes) if font_sizes else None,
                font_name=Counter(font_names).most_common(1)[0][0] if font_names else None,
            )
        )
    return lines


def _infer_pdf_theme(
    lines: list[LayoutLine],
    page_width: float | None,
    page_height: float | None,
) -> DocumentTheme:
    font_sizes = [line.font_size for line in lines if line.font_size]
    font_names = [line.font_name for line in lines if line.font_name]
    body_font_size = None
    heading_font_size = None

    if font_sizes:
        body_font_size = Counter(round(size, 1) for size in font_sizes).most_common(1)[0][0]
        larger_sizes = [size for size in font_sizes if size >= body_font_size + 1]
        heading_font_size = max(larger_sizes) if larger_sizes else body_font_size

    body_font_name = Counter(font_names).most_common(1)[0][0] if font_names else None

    return DocumentTheme(
        source_format="pdf",
        page_width=page_width,
        page_height=page_height,
        body_font_name=body_font_name,
        heading_font_name=body_font_name,
        body_font_size=body_font_size,
        heading_font_size=heading_font_size,
        preserve_original_pdf_look=True,
    )


def _docx_paragraph_font_size(paragraph) -> float | None:
    if paragraph.style and getattr(paragraph.style, "font", None) and paragraph.style.font.size:
        return _length_to_pt(paragraph.style.font.size)
    for run in paragraph.runs:
        if run.font and run.font.size:
            return _length_to_pt(run.font.size)
    return None


def _docx_paragraph_font_name(paragraph) -> str | None:
    if paragraph.style and getattr(paragraph.style, "font", None) and paragraph.style.font.name:
        return paragraph.style.font.name
    for run in paragraph.runs:
        if run.font and run.font.name:
            return run.font.name
    return None


def _length_to_pt(length: Length) -> float:
    return round(float(length.pt), 1)
